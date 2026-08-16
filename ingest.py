# 1) 이 파일은 "문서를 읽어서 벡터 DB에 저장하는 과정"을 담당한다.
# RAG 시스템에서 가장 먼저 하는 단계이며,
# 사용자가 질문할 때 문서를 빠르게 비교 검색할 수 있게 준비해주는 코드다.

# 2) 필요한 라이브러리를 불러온다.
# - Path: 파일 위치를 다루는 도구
# - load_dotenv: .env 파일의 환경 변수 값을 읽어온다.
# - TextLoader: .txt, .md 같은 파일 내용을 읽는 도구
# - PyPDFLoader: .pdf 파일을 읽는 도구
# - HuggingFaceEmbeddings: 문장을 숫자 벡터로 바꾸는 임베딩 모델
# - RecursiveCharacterTextSplitter: 문서를 작은 조각(chunk)으로 나누는 도구
# - Chroma: 벡터 DB를 저장하고 관리하는 도구
import gc
import re
import shutil
import time
import uuid
from pathlib import Path

from dotenv import load_dotenv

import pandas as pd
from docx import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma

try:
    from pyhwp import HwpFile
except ImportError:  # pragma: no cover
    HwpFile = None

# 3) .env 파일 안의 환경 변수를 읽어온다.
# 예: ANTHROPIC_API_KEY 같은 값이 필요할 수 있지만,
# 이 스크립트 자체는 문서를 로딩/임베딩하는 역할이 중심이다.
load_dotenv()

# 4) 문서가 있는 폴더 경로를 지정한다.
# 현재 프로젝트 루트 안에 있는 documents 폴더를 의미한다.
DOCUMENTS_DIR = Path("documents")

# 5) 벡터 DB를 저장할 폴더 이름을 정한다.
# 이후 app.py 에서 이 경로를 읽어 문서를 검색한다.
PERSIST_DIRECTORY = "chroma_db"


def clear_persist_directory(path: Path):
    if not path.exists():
        return

    last_error = None
    for _ in range(20):
        try:
            shutil.rmtree(path)
            return
        except (PermissionError, OSError) as exc:
            last_error = exc
            gc.collect()
            time.sleep(0.25)

    if path.exists():
        for child in sorted(path.rglob("*"), reverse=True):
            try:
                if child.is_file() or child.is_symlink():
                    child.unlink()
                elif child.is_dir():
                    child.rmdir()
            except OSError:
                pass
        try:
            path.rmdir()
            return
        except OSError:
            pass

    if last_error is not None:
        raise last_error


# 6) documents 폴더 안의 모든 문서를 읽어오는 함수
# 이 함수가 하는 일: 파일 하나씩 열어서 내용을 리스트 형태로 준비하는 것
# 반환값: 문서 객체들의 목록

def normalize_text(raw_text: str) -> str:
    if raw_text is None:
        return ""

    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\t", " | ")
    text = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n +", "\n", text)
    text = re.sub(r" +\n", "\n", text)
    return text.strip()


def read_text_with_fallback(file_path: Path) -> str:
    try:
        raw_bytes = file_path.read_bytes()
    except OSError:
        return ""

    encodings = ["utf-8-sig", "utf-8", "cp949", "euc-kr", "cp1252", "latin-1"]
    for encoding in encodings:
        try:
            return normalize_text(raw_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    return normalize_text(raw_bytes.decode("utf-8", errors="replace"))


def extract_hwp_text(file_path: Path) -> str:
    if HwpFile is None:
        return ""

    try:
        hwp = HwpFile(file_path)
        hwp_text = getattr(hwp, "text", "")
        if callable(hwp_text):
            hwp_text = hwp_text()
        return normalize_text(str(hwp_text or ""))
    except Exception:
        return ""


def read_text_from_file(file_path: Path):
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        loader = PyPDFLoader(str(file_path))
        pages = loader.load()
        text = "\n\n".join(
            page.page_content.strip()
            for page in pages
            if page.page_content and page.page_content.strip()
        )
        return normalize_text(text)

    if suffix in {".txt", ".md"}:
        return read_text_with_fallback(file_path)

    if suffix == ".csv":
        try:
            df = pd.read_csv(file_path, encoding="utf-8-sig", engine="python", keep_default_na=False)
            return normalize_text(df.to_string(index=False))
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding="cp949", engine="python", keep_default_na=False)
            return normalize_text(df.to_string(index=False))
        except Exception:
            try:
                df = pd.read_csv(file_path, encoding="utf-8", engine="python", keep_default_na=False, on_bad_lines="skip")
                return normalize_text(df.to_string(index=False))
            except Exception:
                return ""

    if suffix == ".docx":
        doc = Document(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return normalize_text("\n".join(paragraphs))

    if suffix in {".xlsx", ".xls"}:
        df = pd.read_excel(file_path, engine="openpyxl" if suffix == ".xlsx" else None)
        return normalize_text(df.to_string(index=False))

    if suffix in {".odt", ".ods", ".odp"}:
        try:
            if suffix == ".ods":
                df = pd.read_excel(file_path, engine="odf")
                return normalize_text(df.to_string(index=False))
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return normalize_text(f.read())
        except Exception:
            return ""

    if suffix == ".hwp":
        return extract_hwp_text(file_path)

    return ""


def load_documents():
    # 문서들을 담을 빈 리스트를 만든다.
    documents = []
    supported_extensions = {".txt", ".md", ".pdf", ".csv", ".docx", ".xlsx", ".xls", ".odt", ".ods", ".odp", ".hwp"}

    # documents 폴더 안의 모든 파일을 순회한다.
    for file_path in DOCUMENTS_DIR.glob("*"):
        if file_path.suffix.lower() not in supported_extensions:
            continue

        text = read_text_from_file(file_path)
        if not text.strip():
            continue

        document = type("Doc", (), {})()
        document.page_content = text
        document.metadata = {"source": file_path.name}
        documents.append(document)

    return documents


# 7) 문서를 잘게 나누고 임베딩을 만들어서 벡터DB에 저장하는 함수
# 이 함수가 실제로 RAG 준비 작업의 핵심이다.
def create_vector_db(persist_directory: str = PERSIST_DIRECTORY, log_callback=None):
    if log_callback is None:
        log_callback = lambda message: None

    # Windows에서는 현재 사용 중인 Chroma 경로를 삭제하려고 하면 잠금 때문에 실패할 수 있다.
    # 따라서 기존 활성 경로를 직접 삭제하지 않고, 새 고유 폴더를 만든 뒤 그곳에 새 DB를 생성한다.
    base_path = Path(persist_directory)
    new_path = base_path.parent / f"{base_path.name}_{uuid.uuid4().hex[:8]}"
    log_callback(f"[1/5] 기준 경로: {base_path}")
    log_callback(f"[2/5] 새 DB 경로 생성 예정: {new_path}")

    gc.collect()
    if base_path.exists() and str(base_path) != str(Path(PERSIST_DIRECTORY)):
        log_callback(f"[3/5] 이전 DB 경로 정리 시도: {base_path}")
        try:
            clear_persist_directory(base_path)
            log_callback("[3/5] 이전 DB 경로 정리 완료")
        except Exception as exc:
            log_callback(f"[3/5] 이전 DB 경로 정리 실패: {exc}")
            pass

    new_path.mkdir(parents=True, exist_ok=True)
    log_callback(f"[4/5] 새 DB 폴더 생성 완료: {new_path}")

    # 먼저 문서 전체를 읽어온다.
    documents = load_documents()
    log_callback(f"[5/5] 문서 로딩 완료: {len(documents)}개")

    # 읽은 문서가 하나도 없으면 에러를 발생시킨다.
    if not documents:
        raise ValueError("documents 폴더에 지원되는 문서가 없습니다. (.txt, .md, .pdf, .csv, .docx, .xlsx, .xls, .odt, .ods, .odp, .hwp)")

    # 문서를 chunk(조각)으로 나눈다.
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=80
    )
    split_documents = splitter.split_documents(documents)
    log_callback(f"[6/6] 청크 생성 완료: {len(split_documents)}개")

    # 문장 임베딩 모델을 준비한다.
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

    # 새 경로에 Chroma 벡터DB를 저장한다.
    Chroma.from_documents(
        documents=split_documents,
        embedding=embeddings,
        persist_directory=str(new_path),
        collection_name="rag_documents"
    )
    log_callback(f"[완료] Chroma 저장 완료: {new_path}")

    # 가장 최근의 활성 DB 경로를 새 경로로 넘기기 위해 리턴한다.
    return len(documents), len(split_documents), str(new_path)


# 11) 이 파일이 직접 실행될 때만 아래 코드가 작동한다.
# import할 때는 실행되지 않는다.
if __name__ == "__main__":
    # 문서 생성 함수 호출
    document_count, chunk_count, new_db_path = create_vector_db()

    # 결과를 콘솔에 출력한다.
    print(f"원본 문서 수: {document_count}")
    print(f"저장된 문서 조각 수: {chunk_count}")
    print(f"새 DB 경로: {new_db_path}")
    print("벡터 DB 생성이 완료되었습니다.")